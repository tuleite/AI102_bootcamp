import requests
import os
from docx import Document
from bs4 import BeautifulSoup
from langchain_openai.chat_models.azure import AzureChatOpenAI

subscription_key = "your-subscription-key-here"
endpoint = "your-endpoint-here"
location = "region-of-your-service-here(eastus, ...)"

target_language = "pt-br"

# Creating a function to translate string texts
def text_translator(text, target_language):
    path = '/translated'
    constructed_url = endpoint + path
    headers = {
        'subscription_key': subscription_key,
        'subscription_region': location,
        'content_type': application/json,
        'client_traceID': str(os.urandom(16))
    }

    body = [{
        'text': text
    }]

    params = {
        'api-version': "3.0",
        'from': 'en',
        'to': [target_language]
    }

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]['translations'][0]['text']

# Running text translation
text_translator('Once upon a time', 'es-es')

# Creating a function to translate a document, using a nested 'text_translator()' 
def document_translator(path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
        translated_text = text_translator(paragraph.text, target_language)
        full_text.append(translated_text)
    
    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)

    path_translated = path.replace(".docx", f"_{target_language}.docx")
    translated_doc.save(path_translated)
    return path_translated

# Translating a document
input_file = 'your-file-path'
document_translator(input_file)

# Creating a function to translate a article from a URL
def extract_text_from_url(url):
    response = requests.get(url)

    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        for script_or_style in soup(['script', 'style']):
            script_or_style.decompose()
        text = soup.get_text(separator=' ')
        # to clean up the text
        lines = (line.strip() for line in text.splitlines())
        parts = (phrase.strip() for line in lines for phrase in line.split(" "))
        clean_text = '\n'.join(part for part in parts if part)

        return clean_text
    
    else:
        print(f'Failed to fetch the URL. Status code: {response.status_code}')
        return None


client = AzureChatOpenAI(
    azure_endpoint = 'your-azure_openai-endpoint',
    api_key = 'your-azure_openai-api-key',
    api_version = 'api-version',
    deployment_name = "gpt-4o-mini"
    max_retries = 0
)

def translate_article(text, lang):
   
    messages = [
       ("system", "Você atua como tradutor de textos"),
       ("user", f"Traduza: {text} para o idioma {lang} e responda apenas com a tradução no formato markdown")
    ]

    response = client.invoke(messages)
    return response.content
   
# Translating text from a url
url = 'your-url-here'
text = extract_text_from_url(url)
article = translate_article(text, target_language)

print(article)